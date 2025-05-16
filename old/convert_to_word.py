import os
import sys
from pathlib import Path
try:
    import markdown
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing required packages...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-markdown", "python-docx"])
    import markdown
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_file, docx_file):
    """
    Convert a Markdown file to a Word document (.docx)
    """
    print(f"Converting {md_file} to {docx_file}...")
    
    # Read the markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert markdown to HTML
    html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    
    # Create a new Word document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "VISTA3D Documentation"
    doc.core_properties.author = "VISTA3D Team"
    
    # Parse the HTML and add to the document
    # This is a simplified approach - for complex markdown with images, tables, etc.,
    # you might need a more sophisticated parser
    
    lines = md_content.split('\n')
    in_code_block = False
    code_content = []
    
    for line in lines:
        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        # Handle code blocks
        elif line.startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph()
                code_run = p.add_run(code_text)
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
        elif in_code_block:
            code_content.append(line)
        # Handle tables (simplified)
        elif line.startswith('|') and '|' in line[1:]:
            # This is a very basic table handler - would need more work for complex tables
            if not hasattr(convert_md_to_docx, 'table_started'):
                convert_md_to_docx.table_started = True
                # Count columns by splitting by | and removing empty strings
                cols = len([x for x in line.split('|') if x.strip()])
                convert_md_to_docx.table = doc.add_table(rows=1, cols=cols)
                convert_md_to_docx.table.style = 'Table Grid'
                # Add header row
                header_cells = convert_md_to_docx.table.rows[0].cells
                headers = [x.strip() for x in line.split('|') if x.strip()]
                for i, header in enumerate(headers):
                    header_cells[i].text = header
            elif line.startswith('|-'):
                # This is the separator line, skip it
                pass
            else:
                # Add a row to the table
                row_cells = convert_md_to_docx.table.add_row().cells
                cells = [x.strip() for x in line.split('|') if x.strip()]
                for i, cell in enumerate(cells):
                    row_cells[i].text = cell
        # Handle regular paragraphs
        elif line.strip() == '':
            if not in_code_block:
                doc.add_paragraph()
        else:
            if not in_code_block:
                p = doc.add_paragraph()
                # Handle basic formatting (bold, italic)
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # Bold text
                        p.add_run(part).bold = True
                    else:
                        # Handle italic
                        subparts = part.split('*')
                        for j, subpart in enumerate(subparts):
                            if j % 2 == 1:  # Italic text
                                p.add_run(subpart).italic = True
                            else:
                                # Handle inline code
                                code_parts = subpart.split('`')
                                for k, code_part in enumerate(code_parts):
                                    if k % 2 == 1:  # Code text
                                        code_run = p.add_run(code_part)
                                        code_run.font.name = 'Courier New'
                                    else:
                                        p.add_run(code_part)
    
    # Save the document
    doc.save(docx_file)
    print(f"Successfully converted to {docx_file}")

if __name__ == "__main__":
    # Get the path to the markdown file
    md_file = Path("VISTA3D_DOCUMENTATION.md")
    docx_file = md_file.with_suffix('.docx')
    
    # Convert the file
    convert_md_to_docx(md_file, docx_file)
